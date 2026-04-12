# -*- coding: utf-8 -*-
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from collections import OrderedDict
from docx import Document
from docx.oxml.ns import qn, nsmap
from copy import deepcopy
import shutil

DOCX_PATH = r"d:\SoftData\Obsidian\JA-algorithm2\A算法模版.docx"
OUTPUT_DIR = r"d:\SoftData\Obsidian\JA-algorithm2"

WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
OMML_NAMESPACE = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
W14_NAMESPACE = '{http://schemas.microsoft.com/office/word/2010/wordml}'
WP_NAMESPACE = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
A_NAMESPACE = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
R_NAMESPACE = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'

nsmap_custom = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'v': 'urn:schemas-microsoft-com:vml',
    'o': 'urn:schemas-microsoft-com:office:office',
    'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
}

def get_bookmark_name_from_id(doc, bookmark_id):
    return None

def omml_to_latex(element):
    result = []
    for child in element:
        tag = child.tag
        if tag.endswith('}r'):
            for sub in child:
                if sub.tag.endswith('}t'):
                    if sub.text:
                        result.append(sub.text)
                elif sub.tag.endswith('}rPr'):
                    pass
        elif tag.endswith('}t'):
            if child.text:
                result.append(child.text)
        elif tag.endswith('}oMath') or tag.endswith('}oMathPara'):
            result.append(omml_to_latex(child))
        elif tag.endswith('}f'):
            for sub in child:
                if sub.tag.endswith('}t'):
                    if sub.text:
                        result.append(sub.text)
        elif tag.endswith('}num'):
            for sub in child:
                if sub.tag.endswith('}t'):
                    if sub.text:
                        result.append(sub.text)
        elif tag.endswith('}den'):
            for sub in child:
                if sub.tag.endswith('}t'):
                    if sub.text:
                        result.append(sub.text)
        elif tag.endswith('}e'):
            result.append(omml_to_latex(child))
        elif tag.endswith('}sSup'):
            for sub in child:
                if sub.tag.endswith('}e'):
                    result.append(omml_to_latex(sub))
                elif sub.tag.endswith('}sup'):
                    result.append('^{' + omml_to_latex(sub) + '}')
        elif tag.endswith('}sSub'):
            for sub in child:
                if sub.tag.endswith('}e'):
                    result.append(omml_to_latex(sub))
                elif sub.tag.endswith('}sub'):
                    result.append('_{' + omml_to_latex(sub) + '}')
        elif tag.endswith('}sPre'):
            pass
        elif tag.endswith('}d'):
            result.append('(')
            for sub in child:
                if sub.tag.endswith('}e'):
                    result.append(omml_to_latex(sub))
            result.append(')')
        elif tag.endswith('}func'):
            for sub in child:
                if sub.tag.endswith('}fName'):
                    for fn in sub:
                        if fn.tag.endswith('}t'):
                            if fn.text:
                                result.append('\\' + fn.text if fn.text in ['sin', 'cos', 'tan', 'log', 'ln', 'sqrt', 'sum', 'prod', 'lim', 'max', 'min'] else fn.text)
                elif sub.tag.endswith('}e'):
                    result.append(omml_to_latex(sub))
        elif tag.endswith('}acc'):
            for sub in child:
                if sub.tag.endswith('}e'):
                    result.append('\\bar{' + omml_to_latex(sub) + '}')
        elif tag.endswith('}rad'):
            for sub in child:
                if sub.tag.endswith('}e'):
                    result.append('\\sqrt{' + omml_to_latex(sub) + '}')
        elif tag.endswith('}limLow'):
            for sub in child:
                if sub.tag.endswith('}e'):
                    result.append(omml_to_latex(sub))
                elif sub.tag.endswith('}lim'):
                    result.append('_{' + omml_to_latex(sub) + '}')
        elif tag.endswith('}limUpp'):
            for sub in child:
                if sub.tag.endswith('}e'):
                    result.append(omml_to_latex(sub))
                elif sub.tag.endswith('}lim'):
                    result.append('^{' + omml_to_latex(sub) + '}')
        elif tag.endswith('}m'):
            row_results = []
            for mr in child:
                if mr.tag.endswith('}mr'):
                    cell_results = []
                    for e in mr:
                        if e.tag.endswith('}e'):
                            cell_results.append(omml_to_latex(e))
                    row_results.append(' & '.join(cell_results))
            if row_results:
                result.append('\\begin{matrix} ' + ' \\\\ '.join(row_results) + ' \\end{matrix}')
        elif tag.endswith('}nary'):
            nary_chr = ''
            nary_lim_loc = 'undOvr'
            for sub in child:
                if sub.tag.endswith('}naryPr'):
                    for pr in sub:
                        if pr.tag.endswith('}chr'):
                            nary_chr = pr.get(qn('m:val'), '')
                        elif pr.tag.endswith('}limLoc'):
                            nary_lim_loc = pr.get(qn('m:val'), 'undOvr')
                elif sub.tag.endswith('}e'):
                    base = omml_to_latex(sub)
                elif sub.tag.endswith('}sub'):
                    sub_val = omml_to_latex(sub)
                elif sub.tag.endswith('}sup'):
                    sup_val = omml_to_latex(sub)
            if nary_chr == '∫':
                result.append('\\int')
            elif nary_chr == '∑':
                result.append('\\sum')
            elif nary_chr == '∏':
                result.append('\\prod')
            else:
                result.append('\\sum')
            if nary_lim_loc == 'undOvr':
                if 'sub_val' in dir():
                    result.append('_{' + sub_val + '}')
                if 'sup_val' in dir():
                    result.append('^{' + sup_val + '}')
            else:
                if 'sup_val' in dir():
                    result.append('^{' + sup_val + '}')
                if 'sub_val' in dir():
                    result.append('_{' + sub_val + '}')
            if 'base' in dir():
                result.append(' ' + base)
        elif tag.endswith('}borderBox'):
            for sub in child:
                if sub.tag.endswith('}e'):
                    result.append('\\boxed{' + omml_to_latex(sub) + '}')
        elif tag.endswith('}eqArr'):
            eqs = []
            for sub in child:
                if sub.tag.endswith('}e'):
                    eqs.append(omml_to_latex(sub))
            result.append(' = '.join(eqs))
        elif tag.endswith('}bar'):
            bar_pos = 'bot'
            for sub in child:
                if sub.tag.endswith('}barPr'):
                    for pr in sub:
                        if pr.tag.endswith('}pos'):
                            bar_pos = pr.get(qn('m:val'), 'bot')
                elif sub.tag.endswith('}e'):
                    content = omml_to_latex(sub)
                    if bar_pos == 'top':
                        result.append('\\overline{' + content + '}')
                    else:
                        result.append('\\underline{' + content + '}')
        elif tag.endswith('}groupChr'):
            chr_val = '⏞'
            pos = 'bot'
            for sub in child:
                if sub.tag.endswith('}groupChrPr'):
                    for pr in sub:
                        if pr.tag.endswith('}chr'):
                            chr_val = pr.get(qn('m:val'), '⏞')
                        elif pr.tag.endswith('}pos'):
                            pos = pr.get(qn('m:val'), 'bot')
                elif sub.tag.endswith('}e'):
                    content = omml_to_latex(sub)
                    if chr_val == '⏞' and pos == 'top':
                        result.append('\\overbrace{' + content + '}')
                    elif chr_val == '⏴' and pos == 'bot':
                        result.append('\\underbrace{' + content + '}')
                    else:
                        result.append(content)
        elif tag.endswith('}phant'):
            pass
        elif tag.endswith('}ctrl'):
            pass
        else:
            result.append(omml_to_latex(child))
    return ''.join(result)

def extract_text_from_element(element, hyperlink_targets=None, bookmark_names=None):
    if hyperlink_targets is None:
        hyperlink_targets = {}
    if bookmark_names is None:
        bookmark_names = {}
    
    text_parts = []
    
    for child in element:
        tag = child.tag
        
        if tag == qn('w:t') or tag.endswith('}t'):
            if child.text:
                text_parts.append(child.text)
        elif tag == qn('w:tab') or tag.endswith('}tab'):
            text_parts.append('\t')
        elif tag == qn('w:br') or tag.endswith('}br'):
            text_parts.append('\n')
        elif tag == qn('w:cr') or tag.endswith('}cr'):
            text_parts.append('\n')
        elif tag == qn('w:noBreakHyphen') or tag.endswith('}noBreakHyphen'):
            text_parts.append('-')
        elif tag == qn('w:softHyphen') or tag.endswith('}softHyphen'):
            text_parts.append('-')
        elif tag == qn('w:sym') or tag.endswith('}sym'):
            char = child.get(qn('w:char'))
            if char:
                try:
                    text_parts.append(chr(int(char, 16)))
                except:
                    text_parts.append(char)
        elif tag == qn('w:hyperlink') or tag.endswith('}hyperlink'):
            hyperlink_text = extract_text_from_element(child, hyperlink_targets, bookmark_names)
            anchor = child.get(qn('w:anchor'))
            if anchor:
                anchor_name = bookmark_names.get(anchor, anchor)
                text_parts.append(f'[[{anchor_name}|{hyperlink_text}]]')
            else:
                text_parts.append(hyperlink_text)
        elif tag == qn('w:r') or tag.endswith('}r'):
            for sub in child:
                sub_tag = sub.tag
                if sub_tag == qn('w:t') or sub_tag.endswith('}t'):
                    if sub.text:
                        text_parts.append(sub.text)
                elif sub_tag == qn('w:tab') or sub_tag.endswith('}tab'):
                    text_parts.append('\t')
                elif sub_tag == qn('w:br') or sub_tag.endswith('}br'):
                    text_parts.append('\n')
                elif sub_tag == qn('w:cr') or sub_tag.endswith('}cr'):
                    text_parts.append('\n')
                elif sub_tag == qn('w:drawing') or sub_tag.endswith('}drawing'):
                    pass
                elif sub_tag == qn('w:pict') or sub_tag.endswith('}pict'):
                    pass
                elif sub_tag == qn('w:object') or sub_tag.endswith('}object'):
                    pass
        elif tag == qn('w:bookmarkStart') or tag.endswith('}bookmarkStart'):
            pass
        elif tag == qn('w:bookmarkEnd') or tag.endswith('}bookmarkEnd'):
            pass
        elif tag == qn('w:commentRangeStart') or tag.endswith('}commentRangeStart'):
            pass
        elif tag == qn('w:commentRangeEnd') or tag.endswith('}commentRangeEnd'):
            pass
        elif tag == qn('w:oMath') or tag.endswith('}oMath') or tag == qn('m:oMath') or tag == '{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath':
            latex = omml_to_latex(child)
            if latex.strip():
                text_parts.append(f'${latex}$')
        elif tag == qn('w:oMathPara') or tag.endswith('}oMathPara') or tag == qn('m:oMathPara') or tag == '{http://schemas.openxmlformats.org/officeDocument/2006/math}oMathPara':
            latex = omml_to_latex(child)
            if latex.strip():
                text_parts.append(f'\n$$\n{latex}\n$$\n')
        else:
            text_parts.append(extract_text_from_element(child, hyperlink_targets, bookmark_names))
    
    return ''.join(text_parts)

def get_paragraph_text_with_math(para, hyperlink_targets=None, bookmark_names=None):
    if hyperlink_targets is None:
        hyperlink_targets = {}
    if bookmark_names is None:
        bookmark_names = {}
    
    para_elem = para._element
    return extract_text_from_element(para_elem, hyperlink_targets, bookmark_names)

def get_heading_level(para):
    style_name = para.style.name if para.style else ''
    if style_name.startswith('Heading ') or style_name.startswith('标题 '):
        try:
            level = int(style_name.split()[-1])
            return level
        except:
            pass
    if 'Heading' in style_name or '标题' in style_name:
        for i in range(1, 10):
            if str(i) in style_name:
                return i
    style_id = para.style.style_id if para.style and hasattr(para.style, 'style_id') else ''
    if style_id.startswith('Heading') or style_id.startswith('heading'):
        for i in range(1, 10):
            if str(i) in style_id:
                return i
    return None

def is_textbox_paragraph(para):
    para_elem = para._element
    parent = para_elem.getparent()
    if parent is not None:
        parent_tag = parent.tag
        if 'txbxContent' in parent_tag or 'textbox' in parent_tag.lower():
            return True
        grandparent = parent.getparent()
        if grandparent is not None:
            if 'txbxContent' in grandparent.tag or 'textbox' in grandparent.tag.lower():
                return True
    return False

def extract_textboxes_from_doc(doc):
    textboxes = []
    
    for shape in doc.inline_shapes:
        pass
    
    return textboxes

def parse_docx_structure(docx_path):
    doc = Document(docx_path)
    
    hyperlink_targets = {}
    bookmark_names = {}
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            try:
                with z.open('word/_rels/document.xml.rels') as f:
                    rels_tree = ET.parse(f)
                    rels_root = rels_tree.getroot()
                    for rel in rels_root.iter():
                        if 'Relationship' in rel.tag:
                            rel_id = rel.get('Id')
                            target = rel.get('Target')
                            if rel_id and target:
                                hyperlink_targets[rel_id] = target
            except:
                pass
            
            try:
                with z.open('word/document.xml') as f:
                    doc_tree = ET.parse(f)
                    doc_root = doc_tree.getroot()
                    for bookmark in doc_root.iter():
                        if 'bookmarkStart' in bookmark.tag:
                            bookmark_id = bookmark.get(qn('w:id'))
                            bookmark_name = bookmark.get(qn('w:name'))
                            if bookmark_id and bookmark_name:
                                bookmark_names[bookmark_id] = bookmark_name
            except:
                pass
    except:
        pass
    
    elements = []
    textbox_buffer = []
    in_textbox = False
    
    for para in doc.paragraphs:
        para_elem = para._element
        
        is_in_textbox = False
        parent = para_elem.getparent()
        while parent is not None:
            parent_tag = parent.tag
            if 'txbxContent' in parent_tag:
                is_in_textbox = True
                break
            parent = parent.getparent()
        
        if is_in_textbox:
            text = get_paragraph_text_with_math(para, hyperlink_targets, bookmark_names)
            textbox_buffer.append(text)
            in_textbox = True
        else:
            if in_textbox and textbox_buffer:
                textbox_content = '\n'.join(textbox_buffer)
                elements.append({
                    'type': 'textbox',
                    'content': textbox_content
                })
                textbox_buffer = []
                in_textbox = False
            
            heading_level = get_heading_level(para)
            text = get_paragraph_text_with_math(para, hyperlink_targets, bookmark_names)
            
            if heading_level:
                elements.append({
                    'type': 'heading',
                    'level': heading_level,
                    'text': text,
                    'content': text
                })
            else:
                if text.strip():
                    elements.append({
                        'type': 'paragraph',
                        'content': text
                    })
    
    if textbox_buffer:
        textbox_content = '\n'.join(textbox_buffer)
        elements.append({
            'type': 'textbox',
            'content': textbox_content
        })
    
    return elements, bookmark_names

def extract_textboxes_from_xml(docx_path):
    textboxes = []
    
    with zipfile.ZipFile(docx_path, 'r') as z:
        with z.open('word/document.xml') as f:
            tree = ET.parse(f)
            root = tree.getroot()
            
            for elem in root.iter():
                if 'txbxContent' in elem.tag:
                    textbox_paras = []
                    for child in elem.iter():
                        if 'p' in child.tag and child.tag.endswith('}p'):
                            para_text = extract_text_from_element(child)
                            if para_text.strip():
                                textbox_paras.append(para_text)
                    
                    if textbox_paras:
                        textboxes.append({
                            'type': 'textbox',
                            'content': '\n'.join(textbox_paras)
                        })
    
    return textboxes

def parse_docx_full(docx_path):
    doc = Document(docx_path)
    
    hyperlink_targets = {}
    bookmark_names = {}
    
    with zipfile.ZipFile(docx_path, 'r') as z:
        try:
            with z.open('word/_rels/document.xml.rels') as f:
                rels_tree = ET.parse(f)
                rels_root = rels_tree.getroot()
                for rel in rels_root.iter():
                    if 'Relationship' in rel.tag:
                        rel_id = rel.get('Id')
                        target = rel.get('Target')
                        if rel_id and target:
                            hyperlink_targets[rel_id] = target
        except:
            pass
    
    elements = []
    
    with zipfile.ZipFile(docx_path, 'r') as z:
        with z.open('word/document.xml') as f:
            tree = ET.parse(f)
            root = tree.getroot()
            
            for bookmark in root.iter():
                if 'bookmarkStart' in bookmark.tag:
                    bookmark_id = bookmark.get(qn('w:id'))
                    bookmark_name = bookmark.get(qn('w:name'))
                    if bookmark_id and bookmark_name:
                        bookmark_names[bookmark_id] = bookmark_name
            
            body = None
            for child in root:
                if 'body' in child.tag:
                    body = child
                    break
            
            if body is None:
                body = root
            
            def process_element(elem, parent_is_textbox=False):
                tag = elem.tag
                
                if 'txbxContent' in tag:
                    textbox_content = []
                    for child in elem:
                        if 'p' in child.tag and child.tag.endswith('}p'):
                            para_text = extract_text_from_element(child, hyperlink_targets, bookmark_names)
                            if para_text.strip():
                                textbox_content.append(para_text)
                    
                    if textbox_content:
                        elements.append({
                            'type': 'textbox',
                            'content': '\n'.join(textbox_content)
                        })
                    return
                
                if 'p' in tag and tag.endswith('}p'):
                    if parent_is_textbox:
                        return
                    
                    is_heading = False
                    heading_level = None
                    
                    for pPr in elem:
                        if 'pPr' in pPr.tag:
                            for pStyle in pPr:
                                if 'pStyle' in pStyle.tag:
                                    style_val = pStyle.get(qn('w:val'))
                                    if style_val:
                                        if style_val.startswith('Heading') or style_val.startswith('heading'):
                                            is_heading = True
                                            for i in range(1, 10):
                                                if str(i) in style_val:
                                                    heading_level = i
                                                    break
                                        elif '标题' in style_val:
                                            is_heading = True
                                            for i in range(1, 10):
                                                if str(i) in style_val:
                                                    heading_level = i
                                                    break
                    
                    para_text = extract_text_from_element(elem, hyperlink_targets, bookmark_names)
                    
                    if is_heading and heading_level:
                        elements.append({
                            'type': 'heading',
                            'level': heading_level,
                            'text': para_text,
                            'content': para_text
                        })
                    else:
                        if para_text.strip():
                            elements.append({
                                'type': 'paragraph',
                                'content': para_text
                            })
                else:
                    for child in elem:
                        process_element(child, parent_is_textbox)
            
            process_element(body)
    
    return elements, bookmark_names

def build_hierarchy(elements):
    root = {'title': '目录', 'level': 0, 'children': [], 'content': []}
    stack = [root]
    
    for elem in elements:
        if elem['type'] == 'heading':
            level = elem['level']
            title = elem['text']
            
            while len(stack) > level:
                stack.pop()
            
            if len(stack) < level:
                while len(stack) < level:
                    placeholder = {
                        'title': '',
                        'level': len(stack),
                        'children': [],
                        'content': []
                    }
                    stack[-1]['children'].append(placeholder)
                    stack.append(placeholder)
            
            node = {
                'title': title,
                'level': level,
                'children': [],
                'content': []
            }
            stack[-1]['children'].append(node)
            stack.append(node)
        else:
            if stack:
                stack[-1]['content'].append(elem)
    
    return root

def sanitize_filename(name):
    invalid_chars = r'[<>:"/\\|?*]'
    name = re.sub(invalid_chars, '_', name)
    name = name.strip()
    if not name:
        name = 'untitled'
    return name

def generate_obsidian_notes(node, output_dir, parent_link=None, is_root=True):
    if node['title'] and node['title'] != '目录':
        filename = sanitize_filename(node['title']) + '.md'
        filepath = os.path.join(output_dir, filename)
        
        content_lines = []
        
        if parent_link:
            content_lines.append(f"[[{parent_link}]]")
            content_lines.append("")
        
        for child in node['children']:
            if child['title']:
                content_lines.append(f"[[{child['title']}]]")
        
        if node['children']:
            content_lines.append("")
        
        for item in node['content']:
            if item['type'] == 'textbox':
                content_lines.append("```cpp")
                content_lines.append(item['content'])
                content_lines.append("```")
                content_lines.append("")
            else:
                content_lines.append(item['content'])
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content_lines))
        
        print(f"Created: {filename}")
        
        for child in node['children']:
            generate_obsidian_notes(child, output_dir, node['title'], is_root=False)
    else:
        for child in node['children']:
            generate_obsidian_notes(child, output_dir, None, is_root=False)

def convert_inline_latex(text):
    text = re.sub(r'\$([^\$]+)\$', r'$$\1$$', text)
    return text

def process_content(content):
    lines = content.split('\n')
    processed_lines = []
    for line in lines:
        line = convert_inline_latex(line)
        processed_lines.append(line)
    return '\n'.join(processed_lines)

def main():
    print("Parsing Word document...")
    elements, bookmark_names = parse_docx_full(DOCX_PATH)
    
    print(f"Found {len(elements)} elements")
    
    heading_count = sum(1 for e in elements if e['type'] == 'heading')
    textbox_count = sum(1 for e in elements if e['type'] == 'textbox')
    print(f"Headings: {heading_count}, Textboxes: {textbox_count}")
    
    print("\nFirst 20 elements:")
    for i, elem in enumerate(elements[:20]):
        if elem['type'] == 'heading':
            print(f"  [{i}] Heading {elem['level']}: {elem['text'][:50]}...")
        elif elem['type'] == 'textbox':
            print(f"  [{i}] Textbox: {elem['content'][:50]}...")
        else:
            print(f"  [{i}] Paragraph: {elem['content'][:50]}...")
    
    print("\nBuilding hierarchy...")
    root = build_hierarchy(elements)
    
    print(f"Root has {len(root['children'])} top-level children")
    
    print("\nGenerating Obsidian notes...")
    generate_obsidian_notes(root, OUTPUT_DIR)
    
    print("\nDone!")

if __name__ == '__main__':
    main()
